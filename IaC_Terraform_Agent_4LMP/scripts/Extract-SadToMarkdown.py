"""
Extract-SadToMarkdown.py
========================
Extracts text from a SAD (.docx) file and renders it as structured Markdown,
then appends a CPF Module Analysis section that cross-references Azure services
mentioned in the SAD against the CPf-schemas catalog.

Usage:
    python scripts/Extract-SadToMarkdown.py [--sad <path-to.docx>] [--out <output.md>] [--schemas <cpf-schemas-dir>]

Defaults:
    --sad     Auto-discovered: first .docx found in <app-repo>/../arch/
    --out     <app-repo>/../arch/sad-analysis.md
    --schemas <toolkit-repo>/templates/cpf-schemas/

Expected workspace layout (matches the .code-workspace convention):
    <workspace-parent>/
    ├── <app-repo>/          ← run the script from here (IaC code only)
    ├── arch/                ← ../arch  — drop the SAD .docx here
    │   ├── <AppName>-SAD-v<N>.docx
    │   ├── sad-analysis.md  (generated)
    │   └── images/          (generated)
    └── IaC_Terraform_Agent_4LMP/  ← toolkit

Requires:  python-docx  (pip install python-docx)
"""

import argparse
import json
import os
import re
import sys
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

# ---------------------------------------------------------------------------
# Try to import python-docx; give a clear install message if missing
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx is not installed. Run:  pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Azure service → CPF module keyword mappings
# Used to auto-detect azure services in SAD text and propose CPF modules
# ---------------------------------------------------------------------------
AZURE_SERVICE_PATTERNS = [
    # (regex_pattern, cpf_module_name_fragment, friendly_service_name)
    (r"\bfunction\s*app\b|\blazer\b|\bFunctions?\b", "linuxfunctionapp|windowsfunctionapp", "Azure Functions"),
    (r"\bapp\s+service\s+plan\b|\bElastic\s+Premium\b|\bEP1\b|\bEP2\b", "appserviceplan", "App Service Plan"),
    (r"\bblob\s+storage\b|\bstorage\s+account\b|\bS3\b", "storageaccount", "Storage Account"),
    (r"\bpostgresql\b|\bpostgres\b|\bflexible\s+server\b|\bRDS\b", "postgresqlserver|postgresql", "PostgreSQL Flexible Server"),
    (r"\bkey\s+vault\b|\bKMS\b|\bsecrets?\s+manager\b", "keyvault", "Key Vault"),
    (r"\bapplication\s+gateway\b|\bAGW\b|\bAPI\s+Gateway\b|\broute\s*53\b", "applicationgateway", "Application Gateway"),
    (r"\bWAF\b|\bweb\s+application\s+firewall\b", "webapplicationfirewallpolicy", "WAF Policy"),
    (r"\bbastion\b", "bastionhost", "Azure Bastion"),
    (r"\bjump\s*h?ost\b|\bjump\s*box\b", "windowsvirtualmachine", "Windows Virtual Machine"),
    (r"\bprivate\s+endpoint\b|\bPE\b", "privateendpoint", "Private Endpoint"),
    (r"\bprivate\s+dns\b", "privatednszone", "Private DNS Zone"),
    (r"\bnetwork\s+security\s+group\b|\bNSG\b", "networksecuritygroup", "Network Security Group"),
    (r"\bsubnet\b", "subnet", "Subnet"),
    (r"\bvirtual\s+network\b|\bVNet\b|\bVPC\b", "virtualnetwork", "Virtual Network"),
    (r"\bapplication\s+insights\b", "applicationinsights", "Application Insights"),
    (r"\blog\s+analytics\b|\bLAW\b", "loganalyticsworkspace", "Log Analytics Workspace"),
    (r"\bAPIM\b|\bAPI\s+management\b", "apimanagement", "API Management"),
    (r"\bevent\s+grid\b", "eventgrid", "Event Grid"),
    (r"\bservice\s+bus\b", "servicebus", "Service Bus"),
    (r"\bredis\b|\bcache\b", "redis", "Redis Cache"),
    (r"\bcontainer\s+registry\b|\bACR\b|\bECR\b", "containerregistry", "Container Registry"),
    (r"\bAKS\b|\bkubernetes\b", "aks|kubernetes", "AKS"),
    (r"\bfirewall\b(?!\s+policy)", "firewall", "Azure Firewall"),
    (r"\bpublic\s+ip\b|\bpublic\s+IP\b", "publicip", "Public IP"),
    (r"\buser.assigned.identit\b|\bUAI\b|\bmanaged.identit\b", "userassignedidentity", "User Assigned Identity"),
    (r"\bresource\s+group\b|\bRG\b", "resourcegroup", "Resource Group"),
    (r"\bdiagnostic\s+setting\b|\bdiagnostics?\b", "monitordiagnosticsetting", "Monitor Diagnostic Setting"),
    (r"\bapp\s+configuration\b", "appconfiguration", "App Configuration"),
]

# Heading style names → Markdown heading level
HEADING_LEVELS = {
    "heading 1": "#",
    "heading 2": "##",
    "heading 3": "###",
    "heading 4": "####",
    "heading 5": "#####",
    "heading 6": "######",
    "title": "#",
    "subtitle": "##",
}

# Table/list styles to skip (boilerplate)
SKIP_STYLES = {
    "document map",
    "footer",
    "header",
}

# XML namespaces used in DrawingML / WordprocessingML
_XML_NS = {
    "a":  "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r":  "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "v":  "urn:schemas-microsoft-com:vml",
    "pkg": "http://schemas.openxmlformats.org/package/2006/relationships",
}
_R_EMBED = f'{{{_XML_NS["r"]}}}embed'
_R_ID    = f'{{{_XML_NS["r"]}}}id'


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _cell_text(cell) -> str:
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _get_para_image_rids(para_element) -> list[str]:
    """Return all image relationship IDs (r:embed / r:id) found in a paragraph's XML."""
    rids: list[str] = []
    # DrawingML images (modern Word / Microsoft 365)
    for blip in para_element.iter(f'{{{_XML_NS["a"]}}}blip'):
        rid = blip.get(_R_EMBED) or blip.get(_R_ID)
        if rid and rid not in rids:
            rids.append(rid)
    # VML images (legacy Word format)
    for imgdata in para_element.iter("{urn:schemas-microsoft-com:vml}imagedata"):
        rid = imgdata.get(_R_ID) or imgdata.get(_R_EMBED)
        if rid and rid not in rids:
            rids.append(rid)
    return rids


def _get_para_alt_text(para_element) -> str:
    """Extract alt text / title from the first wp:docPr element in a paragraph."""
    for docpr in para_element.iter(f'{{{_XML_NS["wp"]}}}docPr'):
        return docpr.get("descr") or docpr.get("name") or ""
    return ""


def _is_table_of_contents(table) -> bool:
    """Heuristic: a 1-column table where every row looks like a ToC entry."""
    if len(table.columns) > 3:
        return False
    first_texts = [_cell_text(row.cells[0]) for row in table.rows[:5]]
    dots_count = sum(1 for t in first_texts if re.search(r"\.{3,}|\t\d+$|\s\d+$", t))
    return dots_count >= 2


def _table_to_markdown(table) -> str:
    if _is_table_of_contents(table):
        return ""  # Skip ToC tables
    rows = table.rows
    if not rows:
        return ""
    lines = []
    # Header row
    header_cells = [_cell_text(c) for c in rows[0].cells]
    lines.append("| " + " | ".join(header_cells) + " |")
    lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
    for row in rows[1:]:
        cells = [_cell_text(c).replace("\n", " ") for c in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def _paragraph_to_markdown(para) -> str:
    style_name = para.style.name.lower() if para.style else ""
    text = para.text.strip()

    if not text:
        return ""

    # Skip known boilerplate styles
    if any(s in style_name for s in SKIP_STYLES):
        return ""

    # Headings
    for style_key, md_prefix in HEADING_LEVELS.items():
        if style_name.startswith(style_key):
            return f"{md_prefix} {text}"

    # List items
    if "list" in style_name or "bullet" in style_name:
        return f"- {text}"
    if "number" in style_name:
        return f"1. {text}"

    # Bold-only paragraphs (often sub-headings in SAD)
    runs_text = [r.text for r in para.runs if r.text.strip()]
    runs_bold = [r.bold for r in para.runs if r.text.strip()]
    if runs_text and all(runs_bold):
        return f"**{text}**"

    return text


# ---------------------------------------------------------------------------
# Extract docx → list of Markdown lines
# ---------------------------------------------------------------------------

def extract_document(docx_path: Path, image_map: dict | None = None) -> list[str]:
    """
    Convert the docx body to Markdown lines, optionally embedding image
    references for every inline drawing found in the document.

    When image_map is provided (built by extract_images_from_docx), each
    paragraph that contains a drawing emits an image reference BEFORE the
    paragraph text so the agent sees the figure in context.
    """
    doc = Document(str(docx_path))
    lines: list[str] = []

    # We iterate the document body XML to preserve table/paragraph order
    body = doc.element.body
    para_idx = 0
    table_idx = 0

    para_list = doc.paragraphs
    table_list = doc.tables

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            if para_idx < len(para_list):
                para = para_list[para_idx]

                # Emit inline image references before any paragraph text
                if image_map is not None:
                    rids = _get_para_image_rids(para._p)
                    if rids:
                        alt_text = _get_para_alt_text(para._p)
                        for rid in rids:
                            if rid in image_map:
                                info = image_map[rid]
                                # Backfill alt_text into the map if discovered here
                                if alt_text and not info["alt_text"]:
                                    info["alt_text"] = alt_text
                                n     = info["index"]
                                label = info["alt_text"] or f"Figure {n}"
                                if info["type"] == "external":
                                    lc = " *(LucidChart)*" if info["is_lucidchart"] else ""
                                    lines.append(
                                        f"> \U0001f517 **Figure {n}{lc}:** [{label}]({info['url']})"
                                    )
                                else:
                                    lines.append(f"![Figure {n}: {label}](images/{info['filename']})")
                                lines.append("")

                md = _paragraph_to_markdown(para)
                if md:
                    lines.append(md)
                    lines.append("")
                para_idx += 1

        elif tag == "tbl":
            if table_idx < len(table_list):
                md = _table_to_markdown(table_list[table_idx])
                if md:
                    lines.append(md)
                    lines.append("")
                table_idx += 1

    return lines


# ---------------------------------------------------------------------------
# Load CPF catalog
# ---------------------------------------------------------------------------

def load_cpf_catalog(schemas_dir: Path) -> dict:
    catalog_path = schemas_dir / "_catalog.json"
    if not catalog_path.exists():
        return {}
    with open(catalog_path, encoding="utf-8-sig") as f:
        return json.load(f)


def load_cpf_schemas(schemas_dir: Path) -> list[dict]:
    schemas = []
    for p in sorted(schemas_dir.glob("cpf-azure-*.json")):
        try:
            with open(p, encoding="utf-8-sig") as f:
                schemas.append(json.load(f))
        except Exception:
            pass
    return schemas


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def extract_images_from_docx(docx_path: Path, images_dir: Path) -> dict[str, dict]:
    """
    Open docx as a zip archive, read word/_rels/document.xml.rels to discover
    all image relationships, then save each embedded image to images_dir.

    External images (e.g. LucidChart embeds) are recorded with their URL but
    not downloaded — the URL is surfaced in the Markdown for the agent.

    Returns a dict mapping rId → image info:
        {
            "index":        int,       # sequential figure number (1-based)
            "type":         str,       # "embedded" | "external"
            "filename":     str|None,  # e.g. "fig-001.png" (None for external)
            "url":          str|None,  # set for external images
            "is_lucidchart": bool,
            "alt_text":     str,       # populated later from paragraph docPr
        }
    """
    images_dir.mkdir(parents=True, exist_ok=True)
    image_map: dict[str, dict] = {}
    fig_counter = 0

    REL_NS  = _XML_NS["pkg"]
    IMG_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"

    with zipfile.ZipFile(str(docx_path), "r") as zf:
        try:
            rels_xml = zf.read("word/_rels/document.xml.rels")
        except KeyError:
            return {}

        root = ET.fromstring(rels_xml)
        for rel in root.iter(f"{{{REL_NS}}}Relationship"):
            rid       = rel.get("Id", "")
            rel_type  = rel.get("Type", "")
            target    = rel.get("Target", "")
            is_extern = rel.get("TargetMode", "") == "External"

            if "image" not in rel_type.lower():
                continue

            if is_extern:
                # External image URL — LucidChart, SharePoint preview, etc.
                fig_counter += 1
                image_map[rid] = {
                    "index":         fig_counter,
                    "type":          "external",
                    "filename":      None,
                    "url":           target,
                    "is_lucidchart": "lucid" in target.lower(),
                    "alt_text":      "",
                }
            else:
                # Embedded media file inside the zip
                media_path = ("word/" + target) if not target.startswith("word/") else target
                ext = Path(target).suffix.lower() or ".png"
                fig_counter += 1
                fig_name = f"fig-{fig_counter:03d}{ext}"
                dest = images_dir / fig_name
                try:
                    dest.write_bytes(zf.read(media_path.lstrip("/")))
                except KeyError:
                    try:
                        dest.write_bytes(zf.read(target.lstrip("/")))
                    except KeyError:
                        fig_counter -= 1
                        continue
                image_map[rid] = {
                    "index":         fig_counter,
                    "type":          "embedded",
                    "filename":      fig_name,
                    "url":           None,
                    "is_lucidchart": False,
                    "alt_text":      "",
                }

    return image_map


def build_image_manifest_section(image_map: dict[str, dict]) -> list[str]:
    """Generate a Markdown section listing all extracted figures as a table."""
    if not image_map:
        return []
    figures = sorted(image_map.values(), key=lambda x: x["index"])
    lines: list[str] = []
    lines.append("---")
    lines.append("")
    lines.append("# Figure Index")
    lines.append("")
    lines.append(
        "> Auto-extracted by `Extract-SadToMarkdown.py`. "
        "Embedded images are saved to `images/`. "
        "External (LucidChart) diagrams are linked by URL."
    )
    lines.append("")
    lines.append("| # | Type | File / Link | Alt Text |")
    lines.append("|---|---|---|---|")
    for fig in figures:
        n   = fig["index"]
        alt = fig["alt_text"] or ""
        if fig["type"] == "external":
            url = fig["url"]
            lc  = " (LucidChart)" if fig["is_lucidchart"] else " (external)"
            ref = f"[View diagram{lc}]({url})"
        else:
            fname = fig["filename"]
            ref   = f"[{fname}](images/{fname})"
        lines.append(f"| {n} | {fig['type']} | {ref} | {alt} |")
    lines.append("")
    return lines


def save_image_manifest_json(image_map: dict[str, dict], out_path: Path) -> None:
    """Save a JSON manifest of all extracted images for downstream agent consumption."""
    figures = sorted(image_map.values(), key=lambda x: x["index"])
    manifest = {
        "figure_count": len(figures),
        "embedded_count": sum(1 for f in figures if f["type"] == "embedded"),
        "external_count": sum(1 for f in figures if f["type"] == "external"),
        "lucidchart_count": sum(1 for f in figures if f["is_lucidchart"]),
        "figures": figures,
    }
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2, ensure_ascii=False)


# ---------------------------------------------------------------------------
# Analyse text for Azure services and map to CPF modules
# ---------------------------------------------------------------------------

def analyse_services(full_text: str, schemas: list[dict]) -> list[dict]:
    """
    Scan full_text for Azure service mentions.
    Returns list of dicts: {service, pattern_matched, cpf_modules: [...]}
    """
    detected: dict[str, dict] = {}

    for pattern, module_fragment, service_name in AZURE_SERVICE_PATTERNS:
        if re.search(pattern, full_text, re.IGNORECASE):
            if service_name not in detected:
                # Find matching CPF modules
                frags = [f.strip() for f in module_fragment.split("|")]
                matched_modules = []
                for schema in schemas:
                    cpf_id = schema.get("cpf_id", "").lower()
                    if any(frag.lower() in cpf_id for frag in frags):
                        matched_modules.append(schema)
                detected[service_name] = {
                    "service": service_name,
                    "regex": pattern,
                    "cpf_modules": matched_modules,
                }

    return list(detected.values())


# ---------------------------------------------------------------------------
# Build CPF Module Analysis section
# ---------------------------------------------------------------------------

def build_analysis_section(detected: list[dict]) -> list[str]:
    lines: list[str] = []
    lines.append("---")
    lines.append("")
    lines.append("# CPF Module Analysis")
    lines.append("")
    lines.append(
        "> This section is auto-generated by `Extract-SadToMarkdown.py`. "
        "It maps Azure services mentioned in the SAD to the corresponding CPF modules "
        "available in the `cpf-schemas/` catalog."
    )
    lines.append("")

    if not detected:
        lines.append("_No Azure service patterns matched._")
        return lines

    lines.append(f"**{len(detected)} Azure service(s) detected.**")
    lines.append("")

    for item in sorted(detected, key=lambda x: x["service"]):
        service = item["service"]
        modules = item["cpf_modules"]
        lines.append(f"## {service}")
        lines.append("")

        if not modules:
            lines.append("_No matching CPF module found in catalog._")
            lines.append("")
            continue

        for m in modules:
            cpf_id = m.get("cpf_id", "N/A")
            module_name = m.get("module_name", "N/A")
            source = m.get("terraform_source", "N/A")
            required = m.get("inputs", {}).get("required", [])
            optional = list(m.get("inputs", {}).get("optional", {}).keys())
            outputs = list(m.get("outputs", {}).keys())

            lines.append(f"### `{cpf_id}`")
            lines.append("")
            lines.append(f"- **Module name**: `{module_name}`")
            lines.append(f"- **Terraform source**: `{source}`")
            lines.append("")

            if required:
                lines.append("**Required inputs:**")
                lines.append("")
                for r in required:
                    lines.append(f"- `{r}`")
                lines.append("")

            if optional:
                # Only show first 10 optional to keep doc readable
                shown = optional[:10]
                extra = len(optional) - len(shown)
                lines.append("**Optional inputs** (first 10):")
                lines.append("")
                for o in shown:
                    lines.append(f"- `{o}`")
                if extra > 0:
                    lines.append(f"- _...and {extra} more_")
                lines.append("")

            if outputs:
                lines.append("**Outputs:**")
                lines.append("")
                for out_name, out_info in m.get("outputs", {}).items():
                    desc = out_info.get("description", "") if isinstance(out_info, dict) else ""
                    lines.append(f"- `{out_name}`" + (f" – {desc}" if desc else ""))
                lines.append("")

    return lines


# ---------------------------------------------------------------------------
# Entrypoint
# ---------------------------------------------------------------------------

def main():
    # toolkit_root: where the script lives (IaC_Terraform_Agent_4LMP/) — used for schemas only
    toolkit_root = Path(__file__).resolve().parent.parent
    # app_root: the directory the user runs from — their app repo root
    app_root = Path.cwd()

    # Auto-discover the SAD .docx.
    # Looks in <app-repo>/../arch/ (sibling workspace folder, e.g. eidp-uiux/arch).
    # Takes the first .docx found (alphabetically).
    def _find_default_sad() -> str:
        arch_dir = app_root.parent / "arch"
        if arch_dir.is_dir():
            candidates = sorted(arch_dir.glob("*.docx"))
            if candidates:
                return str(candidates[0])
        return ""

    parser = argparse.ArgumentParser(description="Extract SAD .docx to Markdown with CPF module analysis")
    parser.add_argument(
        "--sad",
        default=_find_default_sad(),
        help=(
            "Path to the SAD .docx file. "
            "Defaults to the first .docx found in <app-repo>/../arch/. "
            "Example: --sad ../arch/MyApp-SAD-v1.0.docx"
        ),
    )
    parser.add_argument(
        "--out",
        default=str(app_root.parent / "arch" / "sad-analysis.md"),
        help="Output Markdown file path (default: <app-repo>/../arch/sad-analysis.md)",
    )
    parser.add_argument(
        "--schemas",
        default=str(toolkit_root / "templates" / "cpf-schemas"),
        help="Path to cpf-schemas directory (default: <toolkit>/templates/cpf-schemas/)",
    )
    parser.add_argument(
        "--no-analysis",
        action="store_true",
        help="Skip CPF module analysis section",
    )
    parser.add_argument(
        "--no-images",
        action="store_true",
        help="Skip image extraction (do not save images or generate figure index)",
    )
    args = parser.parse_args()

    if not args.sad:
        print(
            "ERROR: No SAD .docx file specified and none found automatically.\n"
            "       Place your SAD .docx in the sibling arch/ folder:\n"
            "         - <app-repo>/../arch/   (e.g. eidp-uiux/arch/)\n"
            "       Or pass --sad <path> explicitly.",
            file=sys.stderr,
        )
        sys.exit(1)

    sad_path    = Path(args.sad)
    out_path    = Path(args.out)
    schemas_dir = Path(args.schemas)
    images_dir  = out_path.parent / "images"

    if not sad_path.exists():
        print(f"ERROR: SAD file not found: {sad_path}", file=sys.stderr)
        sys.exit(1)

    # --- Image extraction ---------------------------------------------------
    image_map: dict = {}
    if not args.no_images:
        print(f"Extracting images to: {images_dir}")
        image_map = extract_images_from_docx(sad_path, images_dir)
        embedded  = sum(1 for i in image_map.values() if i["type"] == "embedded")
        external  = sum(1 for i in image_map.values() if i["type"] == "external")
        lc_count  = sum(1 for i in image_map.values() if i["is_lucidchart"])
        print(f"  Embedded images : {embedded}")
        print(f"  External images : {external}" + (f"  ({lc_count} LucidChart)" if lc_count else ""))

    # --- Document extraction ------------------------------------------------
    print(f"Extracting: {sad_path.name}")
    md_lines  = extract_document(sad_path, image_map or None)
    full_text = "\n".join(md_lines)

    # --- CPF module analysis ------------------------------------------------
    if not args.no_analysis:
        print(f"Loading CPF schemas from: {schemas_dir}")
        schemas = load_cpf_schemas(schemas_dir)
        print(f"  Loaded {len(schemas)} schemas")
        detected = analyse_services(full_text, schemas)
        print(f"  Detected {len(detected)} Azure service(s)")
        analysis_lines = build_analysis_section(detected)
    else:
        analysis_lines = []

    # --- Image manifest section ---------------------------------------------
    image_manifest_lines = build_image_manifest_section(image_map) if image_map else []

    # Combine: document body + figure index + CPF analysis
    output_lines = md_lines + ["", ""] + image_manifest_lines + ["", ""] + analysis_lines

    # Write Markdown output
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(output_lines))
        f.write("\n")

    # Write JSON image manifest (consumed by /analyse-sad agent)
    if image_map:
        manifest_path = out_path.with_name(out_path.stem + "-images.json")
        save_image_manifest_json(image_map, manifest_path)
        print(f"  Image manifest  : {manifest_path}")

    print(f"Written: {out_path}  ({out_path.stat().st_size:,} bytes)")
    print(f"  Lines in document: {len(md_lines)}")
    if not args.no_images and image_map:
        print(f"  Images extracted: {len(image_map)}")
    if not args.no_analysis:
        print(f"  Azure services detected: {len(detected)}")
        for item in sorted(detected, key=lambda x: x['service']):
            mod_ids = [m.get('cpf_id','?') for m in item['cpf_modules']]
            print(f"    [{item['service']}] → {mod_ids if mod_ids else '(no match)'}")


if __name__ == "__main__":
    main()
